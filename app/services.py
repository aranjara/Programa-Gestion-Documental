from __future__ import annotations

import os
import sys
from datetime import datetime
from pathlib import Path
import csv
import re
import subprocess
import shutil

import openpyxl
from openpyxl.styles import Alignment, Border, Side, Font, PatternFill
from openpyxl.drawing.image import Image as XLImage
from openpyxl.drawing.spreadsheet_drawing import AnchorMarker, OneCellAnchor
from openpyxl.drawing.xdr import XDRPositiveSize2D
from openpyxl.worksheet.properties import PageSetupProperties
from openpyxl.utils.units import pixels_to_EMU
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.database import db_cursor
from app.security import hash_password, verify_password


PROJECT_ROOT = Path(__file__).resolve().parent.parent
LOGO_PATH = PROJECT_ROOT / "activos" / "escudo.png"


DEFAULT_USER = ("admin", "Administrador", "123456", "admin")

FUID_HEADER_DEFAULTS = [
    ("entidad_remitente", ""),
    ("entidad_productora", "ALCALDIA DE RIONEGRO"),
    ("unidad_administrativa", "SECRETARIA DE HACIENDA"),
    ("oficina_productora", "SUBSECRETARIA  DE TESORERIA"),
    ("objeto", ""),
    ("anio", ""),
    ("mes", ""),
    ("dia", ""),
    ("n_transferencia", ""),
]

FUID_DETAIL_FIELDS = [
    "Número de orden",
    "Código",
    "Serie / subserie / asuntos",
    "Nombre unidad documental",
    "Fecha inicial",
    "Fecha final",
    "Físico",
    "Electrónico",
    "Caja",
    "Carpeta",
    "Tomo / legajo / libro",
    "Número de folios",
    "Tipo",
    "Cantidad",
    "Ubicación",
    "Cantidad de documentos electrónicos",
    "Tamaño de documentos electrónicos",
    "Notas",
]

DEFAULT_DETAIL_MAP = [
    ("Número de orden", "field", "Número de orden"),
    ("Código", "field", "Código"),
    ("Serie / subserie / asuntos", "field", "Serie / subserie / asuntos"),
    ("Nombre unidad documental", "field", "Nombre unidad documental"),
    ("Fecha inicial", "field", "Fecha inicial"),
    ("Fecha final", "field", "Fecha final"),
    ("Físico", "field", "Físico"),
    ("Electrónico", "field", "Electrónico"),
    ("Caja", "field", "Caja"),
    ("Carpeta", "field", "Carpeta"),
    ("Tomo / legajo / libro", "field", "Tomo / legajo / libro"),
    ("Número de folios", "field", "Número de folios"),
    ("Tipo", "field", "Tipo"),
    ("Cantidad", "field", "Cantidad"),
    ("Ubicación", "field", "Ubicación"),
    ("Cantidad de documentos electrónicos", "field", "Cantidad de documentos electrónicos"),
    ("Tamaño de documentos electrónicos", "field", "Tamaño de documentos electrónicos"),
    ("Notas", "field", "Notas"),
]

ROTULO_CARPETA_FIELDS = [
    "Fondo",
    "Sección",
    "Subsección",
    "Serie",
    "Subserie",
    "Unidad documental",
    "N° de caja",
    "N° de gaveta",
    "N° de carpeta",
    "N° de folios",
    "Fechas extremas",
    "Observaciones",
]

DEFAULT_ROTULO_CARPETA_MAP = [
    ("Fondo", "field", "FONDO"),
    ("Sección", "field", "SECCIÓN"),
    ("Subsección", "field", "SUBSECCIÓN"),
    ("Serie", "field", "SERIE"),
    ("Subserie", "field", "SUBSERIE"),
    ("Unidad documental", "field", "UNIDAD DOCUMENTAL"),
    ("N° de caja", "field", "CAJA"),
    ("N° de gaveta", "fixed", ""),
    ("N° de carpeta", "field", "CARPETA"),
    ("N° de folios", "field", "Número de folios"),
    ("Fechas extremas", "template", "{Fecha inicial} - {Fecha final}"),
    ("Observaciones", "field", "NOTAS"),
]

ROTULO_CARPETA_CONFIG_DEFAULTS = [
    ("campo_busqueda", "Número de orden"),
    ("titulo", "ALCALDIA DE RIONEGRO"),
    ("fondo", "ALCALDÍA"),
    ("seccion", "SECRETARÍA DE HACIENDA"),
    ("subseccion", "SUBSECRETARÍA DE TESORERÍA"),
    ("serie", "COMPROBANTES CONTABLES"),
    ("subserie", "COMPROBANTES CONTABLES DE EGRESO"),
]



ROTULO_CAJA_CONFIG_DEFAULTS = [
    ("campo_busqueda", "Número de orden"),
    ("campo_caja", "Caja"),
    ("campo_unidad_documental", "Nombre unidad documental"),
    ("campo_fecha_inicial", "Fecha inicial"),
    ("campo_fecha_final", "Fecha final"),
    ("titulo", "ALCALDIA DE RIONEGRO"),
    ("dependencia", "SECRETARIA DE HACIENDA-SUBSECRETARIA DE TESORERIA"),
    ("serie", "COMPROBANTES CONTABLES"),
    ("subserie", "COMPROBANTES CONTABLES DE EGRESO"),
    ("texto_consecutivo", "INICIA CON EL COMPROBANTE DE EGRESO N°"),
    ("texto_correlativo", "FINALIZA CON EL COMPROBANTE DE EGRESO N°"),
    ("observaciones", ""),
]

ROTULO_CAJA_FIELDS = [
    "Serie",
    "Subserie",
    "Consecutivo",
    "Correlativo",
]

DEFAULT_ROTULO_CAJA_MAP = [
    ("Serie", "fixed_config", "serie"),
    ("Subserie", "fixed_config", "subserie"),
    ("Consecutivo", "auto", "consecutivo"),
    ("Correlativo", "auto", "correlativo"),
]


def now_str():

    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def normalize_excel_value(value):
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    return str(value).strip()


def sanitize_column_name(name: str) -> str:
    s = (name or "").strip()
    s = s.replace("\n", " ")
    s = re.sub(r"\s+", " ", s)
    return s


def initialize_database():
    with db_cursor() as (_, cur):
        cur.execute("""
            CREATE TABLE IF NOT EXISTS users(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT NOT NULL UNIQUE,
                full_name TEXT NOT NULL,
                password_hash TEXT NOT NULL,
                role TEXT NOT NULL CHECK(role IN ('admin','normal')),
                active INTEGER NOT NULL DEFAULT 1,
                must_change_password INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL,
                last_login TEXT
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS fields(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                column_name TEXT NOT NULL UNIQUE,
                display_name TEXT NOT NULL,
                visible INTEGER NOT NULL DEFAULT 1,
                display_order INTEGER NOT NULL DEFAULT 0,
                default_value TEXT NOT NULL DEFAULT '',
                created_at TEXT NOT NULL
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS records(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                created_by TEXT NOT NULL,
                active INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            )
        """)
        # Ensure column exists for existing databases
        try:
            cur.execute("ALTER TABLE records ADD COLUMN active INTEGER NOT NULL DEFAULT 1")
        except:
            pass
        cur.execute("""
            CREATE TABLE IF NOT EXISTS record_values(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                record_id INTEGER NOT NULL,
                field_id INTEGER NOT NULL,
                value TEXT NOT NULL DEFAULT '',
                UNIQUE(record_id, field_id),
                FOREIGN KEY(record_id) REFERENCES records(id) ON DELETE CASCADE,
                FOREIGN KEY(field_id) REFERENCES fields(id) ON DELETE CASCADE
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS audit_log(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT NOT NULL,
                action TEXT NOT NULL,
                module TEXT NOT NULL,
                record_id TEXT,
                details TEXT,
                created_at TEXT NOT NULL
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS fuid_header_config(
                config_key TEXT PRIMARY KEY,
                config_value TEXT NOT NULL DEFAULT ''
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS fuid_detail_mapping(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                fuid_field TEXT NOT NULL UNIQUE,
                mapping_type TEXT NOT NULL DEFAULT 'field',
                mapping_value TEXT NOT NULL DEFAULT '',
                display_order INTEGER NOT NULL DEFAULT 0
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS rotulo_carpeta_config(
                config_key TEXT PRIMARY KEY,
                config_value TEXT NOT NULL DEFAULT ''
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS rotulo_carpeta_mapping(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                rotulo_field TEXT NOT NULL UNIQUE,
                mapping_type TEXT NOT NULL DEFAULT 'field',
                mapping_value TEXT NOT NULL DEFAULT '',
                display_order INTEGER NOT NULL DEFAULT 0
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS rotulo_caja_config(
                config_key TEXT PRIMARY KEY,
                config_value TEXT NOT NULL DEFAULT ''
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS rotulo_caja_mapping(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                rotulo_field TEXT NOT NULL UNIQUE,
                mapping_type TEXT NOT NULL DEFAULT 'field',
                mapping_value TEXT NOT NULL DEFAULT '',
                display_order INTEGER NOT NULL DEFAULT 0
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS attachments(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                record_id INTEGER NOT NULL,
                filename TEXT NOT NULL,
                filepath TEXT NOT NULL,
                uploaded_by TEXT NOT NULL,
                uploaded_at TEXT NOT NULL,
                FOREIGN KEY(record_id) REFERENCES records(id) ON DELETE CASCADE
            )
        """)

    seed_admin()
    seed_fuid_defaults()
    seed_rotulo_carpeta_defaults()
    seed_rotulo_caja_defaults()


def seed_admin():
    created = False
    with db_cursor() as (_, cur):
        cur.execute("SELECT COUNT(*) AS total FROM users")
        if cur.fetchone()["total"] == 0:
            username, full_name, password, role = DEFAULT_USER
            cur.execute(
                """
                INSERT INTO users(username, full_name, password_hash, role, active, must_change_password, created_at)
                VALUES (?, ?, ?, ?, 1, 1, ?)
                """,
                (username, full_name, hash_password(password), role, now_str()),
            )
            created = True
    if created:
        log_action("system", "create_default_admin", "users", details="Usuario admin inicial creado")


def seed_fuid_defaults():
    with db_cursor() as (_, cur):
        for key, value in FUID_HEADER_DEFAULTS:
            cur.execute(
                "INSERT OR IGNORE INTO fuid_header_config(config_key, config_value) VALUES (?, ?)",
                (key, value),
            )
        for idx, (field_name, mapping_type, mapping_value) in enumerate(DEFAULT_DETAIL_MAP, start=1):
            cur.execute(
                """
                INSERT OR IGNORE INTO fuid_detail_mapping(fuid_field, mapping_type, mapping_value, display_order)
                VALUES (?, ?, ?, ?)
                """,
                (field_name, mapping_type, mapping_value, idx),
            )



def seed_rotulo_carpeta_defaults():
    with db_cursor() as (_, cur):
        for key, value in ROTULO_CARPETA_CONFIG_DEFAULTS:
            cur.execute(
                "INSERT OR IGNORE INTO rotulo_carpeta_config(config_key, config_value) VALUES (?, ?)",
                (key, value),
            )

        for idx, (field_name, mapping_type, mapping_value) in enumerate(DEFAULT_ROTULO_CARPETA_MAP, start=1):
            cur.execute(
                """
                INSERT OR IGNORE INTO rotulo_carpeta_mapping(rotulo_field, mapping_type, mapping_value, display_order)
                VALUES (?, ?, ?, ?)
                """,
                (field_name, mapping_type, mapping_value, idx),
            )


def seed_rotulo_caja_defaults():
    with db_cursor() as (_, cur):
        for key, value in ROTULO_CAJA_CONFIG_DEFAULTS:
            cur.execute(
                "INSERT OR IGNORE INTO rotulo_caja_config(config_key, config_value) VALUES (?, ?)",
                (key, value),
            )

        for idx, (field_name, mapping_type, mapping_value) in enumerate(DEFAULT_ROTULO_CAJA_MAP, start=1):
            cur.execute(
                """
                INSERT OR IGNORE INTO rotulo_caja_mapping(rotulo_field, mapping_type, mapping_value, display_order)
                VALUES (?, ?, ?, ?)
                """,
                (field_name, mapping_type, mapping_value, idx),
            )


def log_action(username, action, module, record_id=None, details=None):
    with db_cursor() as (_, cur):
        cur.execute(
            """
            INSERT INTO audit_log(username, action, module, record_id, details, created_at)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (username, action, module, record_id, details or "", now_str()),
        )


def authenticate(username, password):
    with db_cursor() as (_, cur):
        cur.execute("SELECT * FROM users WHERE username = ?", (username.strip(),))
        row = cur.fetchone()
        if not row:
            return None, "Usuario no encontrado."
        if int(row["active"]) != 1:
            return None, "Usuario inactivo."
        if not verify_password(password, row["password_hash"]):
            return None, "Clave incorrecta."
        cur.execute("UPDATE users SET last_login = ? WHERE id = ?", (now_str(), row["id"]))
    log_action(username, "login", "auth")
    return dict(row), ""


def change_password(user_id, new_password):
    with db_cursor() as (_, cur):
        cur.execute(
            "UPDATE users SET password_hash = ?, must_change_password = 0 WHERE id = ?",
            (hash_password(new_password), user_id),
        )


def verify_current_password(user_id, current_password):
    with db_cursor() as (_, cur):
        cur.execute("SELECT password_hash FROM users WHERE id = ?", (user_id,))
        row = cur.fetchone()
        return verify_password(current_password, row["password_hash"]) if row else False


def list_users():
    with db_cursor() as (_, cur):
        cur.execute("""
            SELECT id, username, full_name, role, active, must_change_password, created_at, last_login
            FROM users ORDER BY id
        """)
        return [dict(r) for r in cur.fetchall()]


def create_user(current_username, username, full_name, password, role):
    with db_cursor() as (_, cur):
        cur.execute(
            """
            INSERT INTO users(username, full_name, password_hash, role, active, must_change_password, created_at)
            VALUES (?, ?, ?, ?, 1, 1, ?)
            """,
            (username.strip(), full_name.strip(), hash_password(password), role, now_str()),
        )
    log_action(current_username, "create_user", "users", details=f"Usuario creado: {username}")


def update_user_status(current_username, user_id, active):
    with db_cursor() as (_, cur):
        cur.execute("UPDATE users SET active = ? WHERE id = ?", (active, user_id))
    log_action(current_username, "update_user_status", "users", record_id=str(user_id), details=f"Activo={active}")


def reset_user_password(current_username, user_id, temp_password):
    with db_cursor() as (_, cur):
        cur.execute(
            "UPDATE users SET password_hash = ?, must_change_password = 1 WHERE id = ?",
            (hash_password(temp_password), user_id),
        )
    log_action(current_username, "reset_password", "users", record_id=str(user_id))


def update_user_role(current_username, user_id, role):
    with db_cursor() as (_, cur):
        cur.execute("UPDATE users SET role = ? WHERE id = ?", (role, user_id))
    log_action(current_username, "update_user_role", "users", record_id=str(user_id), details=f"Rol={role}")


def get_fields():
    with db_cursor() as (_, cur):
        cur.execute("SELECT * FROM fields ORDER BY display_order, id")
        return [dict(r) for r in cur.fetchall()]


def get_visible_fields():
    return [f for f in get_fields() if int(f["visible"]) == 1]


def get_field_by_column_name(column_name: str):
    with db_cursor() as (_, cur):
        cur.execute("SELECT * FROM fields WHERE column_name = ?", (sanitize_column_name(column_name),))
        row = cur.fetchone()
        return dict(row) if row else None


def create_field(current_username, column_name, display_name, visible, display_order, default_value):
    column_name = sanitize_column_name(column_name)
    display_name = (display_name or "").strip() or column_name
    with db_cursor() as (_, cur):
        cur.execute(
            """
            INSERT INTO fields(column_name, display_name, visible, display_order, default_value, created_at)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (column_name, display_name, int(visible), int(display_order), default_value or "", now_str()),
        )
    log_action(current_username, "create_field", "fields", details=f"Campo creado: {column_name}")


def update_field(current_username, field_id, column_name, display_name, visible, display_order, default_value):
    column_name = sanitize_column_name(column_name)
    display_name = (display_name or "").strip() or column_name
    with db_cursor() as (_, cur):
        cur.execute(
            """
            UPDATE fields
            SET column_name = ?, display_name = ?, visible = ?, display_order = ?, default_value = ?
            WHERE id = ?
            """,
            (column_name, display_name, int(visible), int(display_order), default_value or "", field_id),
        )
    log_action(current_username, "update_field", "fields", record_id=str(field_id), details=f"Campo actualizado: {column_name}")


def delete_field(current_username, field_id):
    with db_cursor() as (_, cur):
        cur.execute("DELETE FROM record_values WHERE field_id = ?", (field_id,))
        cur.execute("DELETE FROM fields WHERE id = ?", (field_id,))
    log_action(current_username, "delete_field", "fields", record_id=str(field_id))


def upsert_record_value(cur, record_id, field_id, value):
    cur.execute(
        """
        INSERT INTO record_values(record_id, field_id, value)
        VALUES (?, ?, ?)
        ON CONFLICT(record_id, field_id) DO UPDATE SET value = excluded.value
        """,
        (record_id, field_id, value or ""),
    )


def create_record(current_username, payload: dict):
    fields = get_fields()
    with db_cursor() as (_, cur):
        cur.execute(
            "INSERT INTO records(created_by, created_at, updated_at) VALUES (?, ?, ?)",
            (current_username, now_str(), now_str()),
        )
        record_id = cur.lastrowid
        for f in fields:
            value = str(payload.get(f["column_name"], "") or "").strip()
            if not value:
                value = f.get("default_value", "") or ""
            upsert_record_value(cur, record_id, f["id"], value)
    log_action(current_username, "create_record", "records", record_id=str(record_id))
    return record_id


def update_record(current_username, record_id, payload: dict):
    fields = get_fields()
    with db_cursor() as (_, cur):
        cur.execute("UPDATE records SET updated_at = ? WHERE id = ?", (now_str(), record_id))
        for f in fields:
            value = str(payload.get(f["column_name"], "") or "").strip()
            if not value:
                value = f.get("default_value", "") or ""
            upsert_record_value(cur, record_id, f["id"], value)
    log_action(current_username, "update_record", "records", record_id=str(record_id))


def delete_record(current_username, record_id):
    with db_cursor() as (_, cur):
        cur.execute("DELETE FROM record_values WHERE record_id = ?", (record_id,))
        cur.execute("DELETE FROM records WHERE id = ?", (record_id,))
    log_action(current_username, "delete_record", "records", record_id=str(record_id))


def update_record_status(current_username, record_id, active):
    with db_cursor() as (_, cur):
        cur.execute("UPDATE records SET active = ?, updated_at = ? WHERE id = ?", (active, now_str(), record_id))
    log_action(current_username, "update_record_status", "records", record_id=str(record_id), details=f"Activo={active}")


def get_record(record_id):
    fields = get_fields()
    field_map = {f["id"]: f for f in fields}
    result = {"id": record_id}
    with db_cursor() as (_, cur):
        cur.execute("SELECT * FROM records WHERE id = ?", (record_id,))
        row = cur.fetchone()
        if not row:
            return None
        result["created_by"] = row["created_by"]
        result["active"] = row["active"]
        result["created_at"] = row["created_at"]
        result["updated_at"] = row["updated_at"]
        cur.execute("SELECT field_id, value FROM record_values WHERE record_id = ?", (record_id,))
        for rv in cur.fetchall():
            field = field_map.get(rv["field_id"])
            if field:
                result[field["column_name"]] = rv["value"]
    return result


def list_records_paged(search_text="", page=1, per_page=20):
    page = max(1, int(page))
    per_page = max(1, int(per_page))
    offset = (page - 1) * per_page
    
    rows_out = []
    total_count = 0
    
    with db_cursor() as (_, cur):
        # We still need to handle search. For simplicity with EAV, we fetch IDs first if search exists
        if search_text.strip():
            q = f"%{search_text.strip()}%"
            # Search in record_values
            cur.execute("""
                SELECT DISTINCT record_id FROM record_values 
                WHERE value LIKE ?
            """, (q,))
            valid_ids = [r["record_id"] for r in cur.fetchall()]
            if not valid_ids:
                return {"records": [], "total": 0, "page": page, "pages": 0}
            
            cur.execute("SELECT COUNT(*) as total FROM records WHERE id IN ({})".format(",".join(map(str, valid_ids))))
            total_count = cur.fetchone()["total"]
            
            cur.execute("""
                SELECT id, created_by, active, created_at, updated_at 
                FROM records 
                WHERE id IN ({})
                ORDER BY id DESC LIMIT ? OFFSET ?
            """.format(",".join(map(str, valid_ids))), (per_page, offset))
        else:
            cur.execute("SELECT COUNT(*) as total FROM records")
            total_count = cur.fetchone()["total"]
            
            cur.execute("""
                SELECT id, created_by, active, created_at, updated_at 
                FROM records 
                ORDER BY id DESC LIMIT ? OFFSET ?
            """, (per_page, offset))
            
        records = cur.fetchall()
        fields = get_fields()
        field_map = {f["id"]: f for f in fields}
        
        for rec in records:
            row = {
                "id": rec["id"],
                "created_by": rec["created_by"],
                "active": rec["active"],
                "created_at": rec["created_at"],
                "updated_at": rec["updated_at"],
            }
            cur.execute("SELECT field_id, value FROM record_values WHERE record_id = ?", (rec["id"],))
            for rv in cur.fetchall():
                f = field_map.get(rv["field_id"])
                if f: row[f["column_name"]] = rv["value"]
            
            # Get attachments count
            cur.execute("SELECT COUNT(*) as total FROM attachments WHERE record_id = ?", (rec["id"],))
            row["attachments_count"] = cur.fetchone()["total"]
            
            rows_out.append(row)
            
    return {
        "records": rows_out,
        "total": total_count,
        "page": page,
        "pages": (total_count + per_page - 1) // per_page
    }


def save_attachment(current_username, record_id, filename, filepath):
    with db_cursor() as (_, cur):
        cur.execute("""
            INSERT INTO attachments(record_id, filename, filepath, uploaded_by, uploaded_at)
            VALUES (?, ?, ?, ?, ?)
        """, (record_id, filename, str(filepath), current_username, now_str()))
    log_action(current_username, "upload_attachment", "records", record_id=str(record_id), details=filename)


def get_attachments(record_id):
    with db_cursor() as (_, cur):
        cur.execute("SELECT * FROM attachments WHERE record_id = ? ORDER BY id DESC", (record_id,))
        return [dict(r) for r in cur.fetchall()]


def delete_attachment(current_username, attachment_id):
    with db_cursor() as (_, cur):
        cur.execute("SELECT filepath FROM attachments WHERE id = ?", (attachment_id,))
        row = cur.fetchone()
        if row and os.path.exists(row["filepath"]):
            try: os.remove(row["filepath"])
            except: pass
        cur.execute("DELETE FROM attachments WHERE id = ?", (attachment_id,))
    log_action(current_username, "delete_attachment", "records", details=f"ID: {attachment_id}")


def list_records(search_text=""):
    rows_out = []
    with db_cursor() as (_, cur):
        cur.execute("SELECT id, created_by, created_at, updated_at FROM records ORDER BY id DESC")
        records = cur.fetchall()
        for rec in records:
            row = {
                "id": rec["id"],
                "created_by": rec["created_by"],
                "created_at": rec["created_at"],
                "updated_at": rec["updated_at"],
            }
            cur.execute(
                """
                SELECT f.column_name, rv.value
                FROM record_values rv
                JOIN fields f ON f.id = rv.field_id
                WHERE rv.record_id = ?
                """,
                (rec["id"],),
            )
            for val in cur.fetchall():
                row[val["column_name"]] = val["value"]
            if search_text.strip():
                q = search_text.strip().lower()
                if not any(q in str(v).lower() for v in row.values()):
                    continue
            rows_out.append(row)
    return rows_out


def create_missing_fields_from_excel_headers(current_username, headers):
    created = 0
    existing = {f["column_name"].lower(): f for f in get_fields()}
    max_order = max([f["display_order"] for f in get_fields()], default=0)
    for header in headers:
        clean = sanitize_column_name(header)
        if not clean:
            continue
        if clean.lower() not in existing:
            max_order += 1
            create_field(current_username, clean, clean, 1, max_order, "")
            existing[clean.lower()] = {"column_name": clean}
            created += 1
    return created


def import_excel_data(current_username, excel_path):
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    if "Datos" not in wb.sheetnames:
        raise ValueError("El archivo no contiene la hoja 'Datos'.")
    ws = wb["Datos"]
    headers = [sanitize_column_name(normalize_excel_value(ws.cell(1, c).value)) for c in range(1, ws.max_column + 1)]
    if not any(headers):
        raise ValueError("La hoja 'Datos' no tiene encabezados válidos.")
    created_fields = create_missing_fields_from_excel_headers(current_username, [h for h in headers if h])
    fields = get_fields()
    field_lookup = {f["column_name"].lower(): f for f in fields}
    imported = 0
    skipped = 0
    for row_num in range(2, ws.max_row + 1):
        row_payload = {}
        has_excel_value = False
        for idx, header in enumerate(headers, start=1):
            if not header:
                continue
            value = normalize_excel_value(ws.cell(row_num, idx).value)
            if value != "":
                has_excel_value = True
            if header.lower() in field_lookup:
                row_payload[field_lookup[header.lower()]["column_name"]] = value
        if not has_excel_value:
            skipped += 1
            continue
        create_record(current_username, row_payload)
        imported += 1
    log_action(
        current_username,
        "import_excel_data",
        "records",
        details=f"Archivo: {excel_path} | Importados: {imported} | Vacíos: {skipped} | Campos creados: {created_fields}",
    )
    return {"imported": imported, "skipped": skipped, "created_fields": created_fields}


def export_records(output_path):
    fields = get_fields()
    rows = list_records()
    headers = ["id"] + [f["column_name"] for f in fields] + ["created_by", "created_at", "updated_at"]
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    if output.suffix.lower() == ".xlsx":
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Datos"
        ws.append(headers)
        for row in rows:
            ws.append([row.get(h, "") for h in headers])
        wb.save(output)
        return str(output)
    csv_output = output.with_suffix(".csv")
    with csv_output.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(headers)
        for row in rows:
            writer.writerow([row.get(h, "") for h in headers])
    return str(csv_output)


def get_fuid_header_config():
    with db_cursor() as (_, cur):
        cur.execute("SELECT config_key, config_value FROM fuid_header_config ORDER BY config_key")
        data = {r["config_key"]: r["config_value"] for r in cur.fetchall()}
    for key, value in FUID_HEADER_DEFAULTS:
        data.setdefault(key, value)
    return data


def save_fuid_header_config(current_username, data: dict):
    with db_cursor() as (_, cur):
        for key, value in data.items():
            cur.execute(
                """
                INSERT INTO fuid_header_config(config_key, config_value)
                VALUES (?, ?)
                ON CONFLICT(config_key) DO UPDATE SET config_value = excluded.config_value
                """,
                (key, value or ""),
            )
    log_action(current_username, "save_fuid_header_config", "fuid")


def get_fuid_detail_mapping():
    with db_cursor() as (_, cur):
        cur.execute("SELECT * FROM fuid_detail_mapping ORDER BY display_order, id")
        rows = [dict(r) for r in cur.fetchall()]
    existing = {r["fuid_field"] for r in rows}
    if len(existing) < len(FUID_DETAIL_FIELDS):
        with db_cursor() as (_, cur):
            for idx, f in enumerate(FUID_DETAIL_FIELDS, start=1):
                if f not in existing:
                    cur.execute(
                        """
                        INSERT INTO fuid_detail_mapping(fuid_field, mapping_type, mapping_value, display_order)
                        VALUES (?, 'field', '', ?)
                        """,
                        (f, idx),
                    )
        return get_fuid_detail_mapping()
    return rows


def update_fuid_detail_mapping(current_username, mapping_id, mapping_type, mapping_value):
    with db_cursor() as (_, cur):
        cur.execute(
            "UPDATE fuid_detail_mapping SET mapping_type = ?, mapping_value = ? WHERE id = ?",
            (mapping_type, mapping_value or "", mapping_id),
        )
    log_action(current_username, "update_fuid_detail_mapping", "fuid", record_id=str(mapping_id))


def resolve_template_expression(template: str, row: dict):
    text = template or ""
    for key, value in row.items():
        text = text.replace("{" + str(key) + "}", str(value or ""))
    return text


def resolve_mapping_value(mapping_row: dict, record_row: dict):
    mapping_type = mapping_row["mapping_type"]
    mapping_value = mapping_row["mapping_value"] or ""

    def buscar_campo(nombre):
        if nombre in record_row:
            return str(record_row.get(nombre, "") or "")

        target = str(nombre).strip().lower()
        for key, value in record_row.items():
            if str(key).strip().lower() == target:
                return str(value or "")

        return None

    if mapping_type == "field":
        encontrado = buscar_campo(mapping_value)
        return encontrado if encontrado is not None else ""

    if mapping_type == "fixed":
        encontrado = buscar_campo(mapping_value)
        return encontrado if encontrado is not None else mapping_value

    if mapping_type == "template":
        return resolve_template_expression(mapping_value, record_row)

    return ""



def get_fuid_records(search_text=""):
    return list_records(search_text)


def _style_cell(cell, border=None, bold=False, align="center", wrap=True, fill=None):
    cell.alignment = Alignment(horizontal=align, vertical="center", wrap_text=wrap)
    if border:
        cell.border = border
    if bold:
        cell.font = Font(bold=True)
    if fill:
        cell.fill = fill


def generate_fuid_excel(current_username, search_text, output_path, template_path=None):
    records = get_fuid_records(search_text)
    if not records:
        raise ValueError("No hay registros para generar el FUID.")

    header = get_fuid_header_config()
    mappings = get_fuid_detail_mapping()

    if output_path is None or str(output_path).strip() == "":
        output_path = PROJECT_ROOT / "salidas" / f"FUID_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "FUID"

    widths = {
        "A": 16, "B": 16, "C": 42, "D": 42, "E": 15, "F": 15, "G": 12, "H": 14,
        "I": 8, "J": 10, "K": 20, "L": 14, "M": 12, "N": 12, "O": 20, "P": 14,
        "Q": 18, "R": 18,
    }
    for col, w in widths.items():
        ws.column_dimensions[col].width = w

    ws.row_dimensions[1].height = 38
    ws.row_dimensions[2].height = 38
    ws.row_dimensions[3].height = 25

    thin = Side(style="thin", color="000000")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    gray = PatternFill("solid", fgColor="D9E1F2")

    try:
        ws.merge_cells("A1:B3")
    except Exception:
        pass

    if LOGO_PATH.exists():
        try:
            img = XLImage(str(LOGO_PATH))
            img.width = 150
            img.height = 95
            ws.add_image(img, "A1")
        except Exception as exc:
            print("No se pudo insertar el escudo:", exc)

    ws["C1"] = "FORMATO UNICO DE INVENTARIO DOCUMENTAL"
    ws.merge_cells("C1:R3")
    _style_cell(ws["C1"], border=border, bold=True, align="center")

    ws["A1"].border = border
    for row in range(1, 4):
        for col in range(3, 19):
            ws.cell(row=row, column=col).border = border

    header_rows = [
        ("A5", "ENTIDAD REMITENTE:", "B5", header.get("entidad_remitente", "")),
        ("A6", "ENTIDAD PRODUCTORA:", "B6", header.get("entidad_productora", "")),
        ("A7", "UNIDAD ADMINISTRATIVA:", "B7", header.get("unidad_administrativa", "")),
        ("A8", "OFICINA PRODUCTORA:", "B8", header.get("oficina_productora", "")),
        ("A9", "OBJETO:", "B9", header.get("objeto", "")),
    ]
    for label_cell, label, value_cell, value in header_rows:
        ws[label_cell] = label
        ws[value_cell] = value
        _style_cell(ws[label_cell], border=border, bold=True, align="left")
        _style_cell(ws[value_cell], border=border, align="left")

    ws["N5"] = "HOJA No.:"
    ws["O5"] = "1"
    ws["P5"] = "DE:"
    ws["Q5"] = "1"
    ws["N6"] = "REGISTRO DE ENTRADA"
    ws.merge_cells("N6:Q6")
    ws["N7"] = "AÑO"
    ws["O7"] = "MES"
    ws["P7"] = "DIA"
    ws["Q7"] = "N° T"
    ws["N8"] = header.get("anio", "")
    ws["O8"] = header.get("mes", "")
    ws["P8"] = header.get("dia", "")
    ws["Q8"] = header.get("n_transferencia", "")

    for row in range(5, 9):
        for col in range(14, 18):
            _style_cell(ws.cell(row=row, column=col), border=border, bold=row in [6, 7], align="center")

    header_row = 12
    data_row = 13

    field_to_col = {
        "Número de orden": 1,
        "Código": 2,
        "Serie / subserie / asuntos": 3,
        "Nombre unidad documental": 4,
        "Fecha inicial": 5,
        "Fecha final": 6,
        "Físico": 7,
        "Electrónico": 8,
        "Caja": 9,
        "Carpeta": 10,
        "Tomo / legajo / libro": 11,
        "Número de folios": 12,
        "Tipo": 13,
        "Cantidad": 14,
        "Ubicación": 15,
        "Cantidad de documentos electrónicos": 16,
        "Tamaño de documentos electrónicos": 17,
        "Notas": 18,
    }

    for m in mappings:
        col = field_to_col.get(m["fuid_field"])
        if col:
            cell = ws.cell(row=header_row, column=col)
            cell.value = m["fuid_field"]
            _style_cell(cell, border=border, bold=True, align="center", fill=gray)

    for idx, record in enumerate(records, start=data_row):
        for m in mappings:
            col = field_to_col.get(m["fuid_field"])
            if not col:
                continue
            value = resolve_mapping_value(m, record)
            cell = ws.cell(row=idx, column=col)
            cell.value = value
            if col in [3, 4, 15, 18]:
                _style_cell(cell, border=border, align="left")
            else:
                _style_cell(cell, border=border, align="center")

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    # Impresión PDF/Excel: hoja horizontal y ajustar a 1 página de ancho.
    ws.page_setup.orientation = "landscape"
    ws.page_setup.paperSize = ws.PAPERSIZE_LEGAL
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr = PageSetupProperties(fitToPage=True)

    ws.page_margins.left = 0.15
    ws.page_margins.right = 0.15
    ws.page_margins.top = 0.25
    ws.page_margins.bottom = 0.255
    ws.page_margins.header = 0.1
    ws.page_margins.footer = 0.1

    ultima_fila = data_row + len(records)
    ws.print_area = f"A1:R{ultima_fila}"
    ws.print_title_rows = f"{header_row}:{header_row}"

    wb.save(out)
    log_action(current_username, "generate_fuid_excel", "fuid", details=str(out))
    return str(out)



def _find_libreoffice():
    posibles_rutas = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]

    for ruta in posibles_rutas:
        if Path(ruta).exists():
            return ruta

    for cmd in ["soffice", "libreoffice"]:
        ruta = shutil.which(cmd)
        if ruta:
            return ruta

    return None


def _find_libreoffice():
    """
    Busca LibreOffice en modo portable, instalado o PATH.
    En el .exe usa la carpeta donde está ArchivoDocumental.exe.
    """
    base = Path(sys.executable).resolve().parent if getattr(sys, "frozen", False) else PROJECT_ROOT

    rutas = [
        base / "LibreOfficePortable" / "App" / "libreoffice" / "program" / "soffice.com",
        base / "LibreOfficePortable" / "App" / "libreoffice" / "program" / "soffice.exe",
        base / "LibreOfficePortable" / "program" / "soffice.com",
        base / "LibreOfficePortable" / "program" / "soffice.exe",
        Path(r"C:\Program Files\LibreOffice\program\soffice.com"),
        Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.com"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
    ]

    for r in rutas:
        if r.exists():
            return str(r)

    return shutil.which("soffice") or shutil.which("libreoffice")


def generate_fuid_pdf_from_excel(excel_path, output_dir):
    """
    Convierte un Excel a PDF usando LibreOffice.
    Compatible con LibreOffice Portable incluido junto al .exe.
    Si falla, crea error_libreoffice.txt en la carpeta de salida.
    """
    output_dir = Path(output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    libreoffice = _find_libreoffice()
    if not libreoffice:
        raise RuntimeError("No se encontró LibreOffice. Verifica LibreOfficePortable o instala LibreOffice.")

    excel_path = Path(excel_path).resolve()

    lo_program_dir = Path(libreoffice).resolve().parent
    profile_dir = output_dir / "lo_profile"
    profile_dir.mkdir(parents=True, exist_ok=True)

    cmd = [
        str(libreoffice),
        "--headless",
        "--invisible",
        "--nologo",
        "--nodefault",
        "--nofirststartwizard",
        "--nolockcheck",
        "--norestore",
        f"-env:UserInstallation={profile_dir.as_uri()}",
        "--convert-to",
        "pdf:calc_pdf_Export",
        "--outdir",
        str(output_dir),
        str(excel_path),
    ]

    env = os.environ.copy()
    env["PATH"] = str(lo_program_dir) + os.pathsep + env.get("PATH", "")

    result = subprocess.run(
        cmd,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        cwd=str(output_dir),
        env=env,
    )

    pdf_path = output_dir / (excel_path.stem + ".pdf")

    if result.returncode != 0 or not pdf_path.exists():
        log_path = output_dir / "error_libreoffice.txt"
        log_path.write_text(
            "COMANDO:\n"
            + " ".join(cmd)
            + "\n\nSTDOUT:\n"
            + str(result.stdout)
            + "\n\nSTDERR:\n"
            + str(result.stderr),
            encoding="utf-8",
        )
        raise RuntimeError(f"No se pudo convertir a PDF con LibreOffice. Revisa: {log_path}")

    return str(pdf_path)



def generate_fuid_word(current_username, search_text, output_path):
    records = get_fuid_records(search_text)
    if not records:
        raise ValueError("No hay registros para generar el FUID.")

    header = get_fuid_header_config()
    mappings = get_fuid_detail_mapping()

    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.25)
    section.bottom_margin = Inches(0.25)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(6)

    header_table = doc.add_table(rows=1, cols=2)
    header_table.style = "Table Grid"
    header_table.autofit = True

    left_cell = header_table.cell(0, 0)
    right_cell = header_table.cell(0, 1)

    if LOGO_PATH.exists():
        try:
            p_logo = left_cell.paragraphs[0]
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_logo = p_logo.add_run()
            run_logo.add_picture(str(LOGO_PATH), width=Inches(1.25))
        except Exception:
            left_cell.text = ""

    p_title = right_cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = p_title.add_run("FORMATO UNICO DE INVENTARIO DOCUMENTAL")
    title_run.bold = True
    title_run.font.size = Pt(11)

    doc.add_paragraph("")

    info_table = doc.add_table(rows=5, cols=4)
    info_table.style = "Table Grid"

    info_rows = [
        ("ENTIDAD REMITENTE:", header.get("entidad_remitente", ""), "HOJA No.:", "1 DE 1"),
        ("ENTIDAD PRODUCTORA:", header.get("entidad_productora", ""), "REGISTRO DE ENTRADA:", ""),
        ("UNIDAD ADMINISTRATIVA:", header.get("unidad_administrativa", ""), "AÑO:", header.get("anio", "")),
        ("OFICINA PRODUCTORA:", header.get("oficina_productora", ""), "MES / DÍA:", f"{header.get('mes', '')} / {header.get('dia', '')}"),
        ("OBJETO:", header.get("objeto", ""), "N° T:", header.get("n_transferencia", "")),
    ]

    for i, row in enumerate(info_rows):
        for j, value in enumerate(row):
            cell = info_table.cell(i, j)
            cell.text = str(value or "")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(7)
                    if j in [0, 2]:
                        run.bold = True

    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=len(mappings))
    table.style = "Table Grid"
    table.autofit = True

    header_cells = table.rows[0].cells
    for i, m in enumerate(mappings):
        header_cells[i].text = m["fuid_field"]
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(5)
                run.bold = True

    text_left_fields = {
        "Serie / subserie / asuntos",
        "Nombre unidad documental",
        "Ubicación",
        "Notas",
    }

    for rec in records:
        row_cells = table.add_row().cells
        for i, m in enumerate(mappings):
            row_cells[i].text = str(resolve_mapping_value(m, rec))
            for paragraph in row_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if m["fuid_field"] in text_left_fields else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(5)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)

    log_action(current_username, "generate_fuid_word", "fuid", details=str(out))
    return str(out)



def generate_fuid(current_username, search_text, output_format, output_dir, template_path=None):
    if output_dir is None or str(output_dir).strip() == "":
        output_dir = str(PROJECT_ROOT / "salidas")
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"FUID_{timestamp}"
    if output_format == "Excel":
        out = output_dir / f"{base_name}.xlsx"
        return generate_fuid_excel(current_username, search_text, out, None)
    if output_format == "PDF":
        xlsx_path = output_dir / f"{base_name}.xlsx"
        generate_fuid_excel(current_username, search_text, xlsx_path, None)
        return generate_fuid_pdf_from_excel(xlsx_path, output_dir)
    if output_format == "Word":
        out = output_dir / f"{base_name}.docx"
        return generate_fuid_word(current_username, search_text, out)
    raise ValueError("Formato de salida no soportado.")



def get_rotulo_carpeta_config():
    with db_cursor() as (_, cur):
        cur.execute("SELECT config_key, config_value FROM rotulo_carpeta_config ORDER BY config_key")
        data = {r["config_key"]: r["config_value"] for r in cur.fetchall()}

    for key, value in ROTULO_CARPETA_CONFIG_DEFAULTS:
        data.setdefault(key, value)

    return data


def save_rotulo_carpeta_config(current_username, data: dict):
    with db_cursor() as (_, cur):
        for key, value in data.items():
            cur.execute(
                """
                INSERT INTO rotulo_carpeta_config(config_key, config_value)
                VALUES (?, ?)
                ON CONFLICT(config_key) DO UPDATE SET config_value = excluded.config_value
                """,
                (key, value or ""),
            )

    log_action(current_username, "save_rotulo_carpeta_config", "rotulo_carpeta")


def get_rotulo_carpeta_mapping():
    """
    Devuelve solo los campos que se configuran desde datos de la base.
    Fondo, Sección, Subsección, Serie y Subserie se manejan como texto fijo
    en Encabezado Rótulo Carpeta.
    """
    excluded = {"Fondo", "Sección", "Subsección", "Serie", "Subserie"}

    with db_cursor() as (_, cur):
        cur.execute("SELECT * FROM rotulo_carpeta_mapping ORDER BY display_order, id")
        rows = [dict(r) for r in cur.fetchall()]

    existing = {r["rotulo_field"] for r in rows}
    if len(existing) < len(ROTULO_CARPETA_FIELDS):
        with db_cursor() as (_, cur):
            for idx, f in enumerate(ROTULO_CARPETA_FIELDS, start=1):
                if f not in existing:
                    cur.execute(
                        """
                        INSERT INTO rotulo_carpeta_mapping(rotulo_field, mapping_type, mapping_value, display_order)
                        VALUES (?, 'field', '', ?)
                        """,
                        (f, idx),
                    )
        return get_rotulo_carpeta_mapping()

    return [r for r in rows if r["rotulo_field"] not in excluded]

def update_rotulo_carpeta_mapping(current_username, mapping_id, mapping_type, mapping_value):
    with db_cursor() as (_, cur):
        cur.execute(
            """
            UPDATE rotulo_carpeta_mapping
            SET mapping_type = ?, mapping_value = ?
            WHERE id = ?
            """,
            (mapping_type, mapping_value or "", mapping_id),
        )

    log_action(current_username, "update_rotulo_carpeta_mapping", "rotulo_carpeta", record_id=str(mapping_id))


def _find_record_value_ci(record_row: dict, field_name: str):
    if field_name in record_row:
        return record_row.get(field_name, "")

    target = str(field_name).strip().lower()
    for key, value in record_row.items():
        if str(key).strip().lower() == target:
            return value

    return ""


def _is_number_like(value):
    try:
        float(str(value).strip().replace(",", "."))
        return True
    except Exception:
        return False


def get_rotulo_carpeta_records(desde, hasta):
    cfg = get_rotulo_carpeta_config()
    campo_busqueda = cfg.get("campo_busqueda", "Número de orden")
    rows = list_records()
    desde_s = str(desde or "").strip()
    hasta_s = str(hasta or desde_s).strip()

    if not desde_s and not hasta_s:
        return rows

    if not hasta_s:
        hasta_s = desde_s

    selected = []
    numeric_range = _is_number_like(desde_s) and _is_number_like(hasta_s)

    if numeric_range:
        ini = float(desde_s.replace(",", "."))
        fin = float(hasta_s.replace(",", "."))
        if ini > fin:
            ini, fin = fin, ini

        for row in rows:
            raw = _find_record_value_ci(row, campo_busqueda)
            if _is_number_like(raw):
                val = float(str(raw).strip().replace(",", "."))
                if ini <= val <= fin:
                    selected.append(row)

        selected.sort(key=lambda r: float(str(_find_record_value_ci(r, campo_busqueda)).strip().replace(",", ".")) if _is_number_like(_find_record_value_ci(r, campo_busqueda)) else 0)
        return selected

    for row in rows:
        raw = str(_find_record_value_ci(row, campo_busqueda)).strip()
        if desde_s == hasta_s:
            if raw == desde_s:
                selected.append(row)
        else:
            if desde_s <= raw <= hasta_s:
                selected.append(row)

    return selected


def _safe_sheet_name(name, fallback):
    raw = str(name or fallback).strip()
    raw = re.sub(r'[\[\]\:\*\?\/\\]', "-", raw)
    raw = raw[:31]
    return raw or fallback


def _apply_rotulo_page_setup(ws, total_rows=47):
    ws.page_setup.orientation = "portrait"
    ws.page_setup.paperSize = ws.PAPERSIZE_LETTER
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.sheet_properties.pageSetUpPr = PageSetupProperties(fitToPage=True)

    ws.page_margins.left = 0.25
    ws.page_margins.right = 0.255
    ws.page_margins.top = 0.25
    ws.page_margins.bottom = 0.255
    ws.page_margins.header = 0.1
    ws.page_margins.footer = 0.1

    ws.print_area = f"A1:K{total_rows}"


def _apply_outer_cut_border(ws, start_row, end_row, start_col=2, end_col=11):
    """
    Marco exterior de corte separado del formato interno.
    Se dibuja por fuera como guía para tijera, sin invadir las líneas internas.
    Referencia primer rótulo: columnas B:K y filas 2:23.
    """
    cut_side = Side(style="medium", color="000000")

    for col in range(start_col, end_col + 1):
        top = ws.cell(row=start_row, column=col)
        bottom = ws.cell(row=end_row, column=col)

        top.border = Border(
            left=top.border.left,
            right=top.border.right,
            top=cut_side,
            bottom=top.border.bottom,
        )
        bottom.border = Border(
            left=bottom.border.left,
            right=bottom.border.right,
            top=bottom.border.top,
            bottom=cut_side,
        )

    for row in range(start_row, end_row + 1):
        left = ws.cell(row=row, column=start_col)
        right = ws.cell(row=row, column=end_col)

        left.border = Border(
            left=cut_side,
            right=left.border.right,
            top=left.border.top,
            bottom=left.border.bottom,
        )
        right.border = Border(
            left=right.border.left,
            right=cut_side,
            top=right.border.top,
            bottom=right.border.bottom,
        )


def _merge(ws, start_row, start_col, end_row, end_col):
    ws.merge_cells(start_row=start_row, start_column=start_col, end_row=end_row, end_column=end_col)


def _find_first_existing_value(record, candidates):
    for candidate in candidates:
        if candidate in record and str(record.get(candidate, "") or "").strip():
            return str(record.get(candidate, "") or "").strip()

        target = candidate.strip().lower()
        for key, value in record.items():
            if str(key).strip().lower() == target and str(value or "").strip():
                return str(value or "").strip()

    return ""


def _build_single_rotulo_carpeta(ws, record, start_row):
    cfg = get_rotulo_carpeta_config()
    mappings = get_rotulo_carpeta_mapping()

    values = {
        m["rotulo_field"]: resolve_mapping_value(
            {"mapping_type": m["mapping_type"], "mapping_value": m["mapping_value"]},
            record
        )
        for m in mappings
    }

    thin = Side(style="thin", color="000000")
    medium = Side(style="medium", color="000000")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    border_medium = Border(left=medium, right=medium, top=medium, bottom=medium)
    gray = PatternFill("solid", fgColor="EDEDED")

    # Formato interno dentro del marco de corte:
    # corte = B:K, formato = C:J
    c0 = 3  # columna C

    for row in range(start_row, start_row + 23):
        ws.row_dimensions[row].height = 16

    ws.row_dimensions[start_row + 1].height = 24
    ws.row_dimensions[start_row + 2].height = 24
    ws.row_dimensions[start_row + 3].height = 18
    ws.row_dimensions[start_row + 9].height = 24
    ws.row_dimensions[start_row + 10].height = 58
    ws.row_dimensions[start_row + 17].height = 18
    ws.row_dimensions[start_row + 18].height = 38
    ws.row_dimensions[start_row + 19].height = 28

    # Encabezado: logo + título
    _merge(ws, start_row + 1, c0, start_row + 3, c0 + 1)

    if LOGO_PATH.exists():
        try:
            img = XLImage(str(LOGO_PATH))
            img.width = 85
            img.height = 60

            # Centrado visual dentro del bloque combinado C:D.
            # c0 es la columna C en índice 1-based; AnchorMarker usa índice 0-based.
            marker = AnchorMarker(
                col=c0 - 1,
                colOff=pixels_to_EMU(45),
                row=start_row,
                rowOff=pixels_to_EMU(6),
            )

            img.anchor = OneCellAnchor(
                _from=marker,
                ext=XDRPositiveSize2D(
                    pixels_to_EMU(img.width),
                    pixels_to_EMU(img.height)
                )
            )

            ws.add_image(img)

        except Exception as exc:
            print("No se pudo insertar el escudo:", exc)

    titulo = cfg.get("titulo", "ALCALDIA DE RIONEGRO")
    _merge(ws, start_row + 1, c0 + 2, start_row + 3, c0 + 7)
    ws.cell(row=start_row + 1, column=c0 + 2).value = titulo
    _style_cell(ws.cell(row=start_row + 1, column=c0 + 2), border=border, bold=True, align="center")

    # Bordes internos del encabezado: delgados, no marco de corte.
    for row in range(start_row + 1, start_row + 4):
        for col in range(c0, c0 + 8):
            ws.cell(row=row, column=col).border = border

    # Datos fijos y unidad documental
    main_rows = [
        (start_row + 5, "FONDO", cfg.get("fondo", "")),
        (start_row + 6, "SECCIÓN", cfg.get("seccion", "")),
        (start_row + 7, "SUBSECCIÓN", cfg.get("subseccion", "")),
        (start_row + 8, "SERIE", cfg.get("serie", "")),
        (start_row + 9, "SUBSERIE", cfg.get("subserie", "")),
        (start_row + 10, "UNIDAD DOCUMENTAL", values.get("Unidad documental", "")),
    ]

    for row, label, value in main_rows:
        _merge(ws, row, c0, row, c0 + 1)
        _merge(ws, row, c0 + 2, row, c0 + 7)

        ws.cell(row=row, column=c0).value = label
        ws.cell(row=row, column=c0 + 2).value = value

        _style_cell(ws.cell(row=row, column=c0), border=border, bold=True, align="left", fill=gray)
        _style_cell(ws.cell(row=row, column=c0 + 2), border=border, align="left", wrap=True)

        for col in list(range(c0 + 1, c0 + 2)) + list(range(c0 + 3, c0 + 8)):
            ws.cell(row=row, column=col).border = border

    # Caja / gaveta / carpeta / folios
    sub_header_row = start_row + 12
    sub_value_row = start_row + 13

    sub_labels = [
        (c0, c0 + 1, "N° DE CAJA", values.get("N° de caja", "")),
        (c0 + 2, c0 + 3, "N° DE GAVETA", values.get("N° de gaveta", "")),
        (c0 + 4, c0 + 5, "N° DE CARPETA", values.get("N° de carpeta", "")),
        (c0 + 6, c0 + 7, "N° DE FOLIOS", values.get("N° de folios", "")),
    ]

    for c1, c2, label, value in sub_labels:
        _merge(ws, sub_header_row, c1, sub_header_row, c2)
        _merge(ws, sub_value_row, c1, sub_value_row, c2)

        ws.cell(row=sub_header_row, column=c1).value = label
        ws.cell(row=sub_value_row, column=c1).value = value

        _style_cell(ws.cell(row=sub_header_row, column=c1), border=border, bold=True, align="center", fill=gray)
        _style_cell(ws.cell(row=sub_value_row, column=c1), border=border, align="center")

        for col in range(c1, c2 + 1):
            ws.cell(row=sub_header_row, column=col).border = border
            ws.cell(row=sub_value_row, column=col).border = border

    # Fechas extremas, sin línea/columna intermedia marcada.
    fecha_row = start_row + 15

    fecha_inicial = _find_first_existing_value(record, ["Fecha inicial", "FECHA INICIAL", "fecha inicial"])
    fecha_final = _find_first_existing_value(record, ["Fecha final", "FECHA FINAL", "fecha final"])

    fecha_completa = values.get("Fechas extremas", "")
    if not fecha_inicial and not fecha_final and fecha_completa:
        partes = str(fecha_completa).split("-")
        if len(partes) >= 2:
            fecha_inicial = partes[0].strip()
            fecha_final = "-".join(partes[1:]).strip()
        else:
            fecha_inicial = fecha_completa

    _merge(ws, fecha_row, c0, fecha_row, c0 + 1)
    _merge(ws, fecha_row, c0 + 2, fecha_row, c0 + 4)
    _merge(ws, fecha_row, c0 + 5, fecha_row, c0 + 7)

    ws.cell(row=fecha_row, column=c0).value = "FECHAS EXTREMAS"
    ws.cell(row=fecha_row, column=c0 + 2).value = fecha_inicial
    ws.cell(row=fecha_row, column=c0 + 5).value = fecha_final

    _style_cell(ws.cell(row=fecha_row, column=c0), border=border, bold=True, align="left", fill=gray)
    _style_cell(ws.cell(row=fecha_row, column=c0 + 2), border=border, align="center", wrap=True)
    _style_cell(ws.cell(row=fecha_row, column=c0 + 5), border=border, align="center", wrap=True)

    for col in [c0 + 1, c0 + 3, c0 + 4, c0 + 6, c0 + 7]:
        ws.cell(row=fecha_row, column=col).border = border

    # Observaciones: título arriba, contenido abajo usando todo el ancho.
    obs_label_row = start_row + 17
    obs_value_row = start_row + 18

    _merge(ws, obs_label_row, c0, obs_label_row, c0 + 7)
    ws.cell(row=obs_label_row, column=c0).value = "OBSERVACIONES:"
    _style_cell(ws.cell(row=obs_label_row, column=c0), border=border, bold=True, align="left", fill=gray)

    _merge(ws, obs_value_row, c0, obs_value_row + 1, c0 + 7)
    ws.cell(row=obs_value_row, column=c0).value = values.get("Observaciones", "")
    _style_cell(ws.cell(row=obs_value_row, column=c0), border=border, align="left", wrap=True)

    for row in range(obs_label_row, obs_value_row + 2):
        for col in range(c0, c0 + 8):
            ws.cell(row=row, column=col).border = border

    # Marco exterior separado para corte.
    # Se aplica al final para que no lo sobreescriban los bordes internos.
    _apply_outer_cut_border(ws, start_row, start_row + 21, 2, 11)

    # No se agrega fila de número de orden ni líneas sobrantes.

def _build_rotulo_carpeta_sheet(ws, records_pair, sheet_index):
    # Columnas: B:K marco de corte; C:J formato interno.
    widths = {
        "A": 2,
        "B": 4.5,
        "C": 12, "D": 12, "E": 12, "F": 12, "G": 12, "H": 12, "I": 12, "J": 12,
        "K": 4.5,
    }
    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    if len(records_pair) >= 1:
        _build_single_rotulo_carpeta(ws, records_pair[0], 2)

    if len(records_pair) >= 2:
        _build_single_rotulo_carpeta(ws, records_pair[1], 25)

    _apply_rotulo_page_setup(ws, total_rows=47)



def generate_rotulo_carpeta_excel(current_username, desde, hasta, output_path):
    records = get_rotulo_carpeta_records(desde, hasta)
    if not records:
        raise ValueError("No se encontraron registros para generar el Rótulo de Carpeta.")

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    used_names = set()
    sheet_index = 1

    for i in range(0, len(records), 2):
        pair = records[i:i + 2]
        base_name = _safe_sheet_name(f"Rotulos {sheet_index}", f"Rotulos {sheet_index}")
        sheet_name = base_name
        count = 2

        while sheet_name in used_names:
            suffix = f"_{count}"
            sheet_name = (base_name[:31 - len(suffix)] + suffix)[:31]
            count += 1

        used_names.add(sheet_name)
        ws = wb.create_sheet(sheet_name)
        _build_rotulo_carpeta_sheet(ws, pair, sheet_index)
        sheet_index += 1

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out)

    log_action(current_username, "generate_rotulo_carpeta_excel", "rotulo_carpeta", details=str(out))
    return str(out)



def generate_rotulo_carpeta_pdf_from_excel(excel_path, output_dir):
    return generate_fuid_pdf_from_excel(excel_path, output_dir)


def _add_rotulo_carpeta_word_label(doc, record):
    cfg = get_rotulo_carpeta_config()
    mappings = get_rotulo_carpeta_mapping()

    values = {
        m["rotulo_field"]: resolve_mapping_value(
            {"mapping_type": m["mapping_type"], "mapping_value": m["mapping_value"]},
            record
        )
        for m in mappings
    }

    fecha_inicial = _find_first_existing_value(record, ["Fecha inicial", "FECHA INICIAL", "fecha inicial"])
    fecha_final = _find_first_existing_value(record, ["Fecha final", "FECHA FINAL", "fecha final"])

    table = doc.add_table(rows=10, cols=2)
    table.style = "Table Grid"

    if LOGO_PATH.exists():
        try:
            p = table.cell(0, 0).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(LOGO_PATH), width=Inches(0.9))
        except Exception:
            pass

    title_cell = table.cell(0, 1)
    title_cell.text = cfg.get("titulo", "ALCALDIA DE RIONEGRO")

    for paragraph in title_cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

    rows = [
        ("FONDO", cfg.get("fondo", "")),
        ("SECCIÓN", cfg.get("seccion", "")),
        ("SUBSECCIÓN", cfg.get("subseccion", "")),
        ("SERIE", cfg.get("serie", "")),
        ("SUBSERIE", cfg.get("subserie", "")),
        ("UNIDAD DOCUMENTAL", values.get("Unidad documental", "")),
        ("N° DE CAJA / GAVETA / CARPETA / FOLIOS", f"{values.get('N° de caja', '')} / {values.get('N° de gaveta', '')} / {values.get('N° de carpeta', '')} / {values.get('N° de folios', '')}"),
        ("FECHAS EXTREMAS", f"{fecha_inicial}                {fecha_final}"),
        ("OBSERVACIONES", values.get("Observaciones", "")),
    ]

    for i, (label, value) in enumerate(rows, start=1):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = str(value or "")

        for c in [0, 1]:
            for paragraph in table.cell(i, c).paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if c == 1 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(7)
                    if c == 0:
                        run.bold = True

    doc.add_paragraph("")


def generate_rotulo_carpeta_word(current_username, desde, hasta, output_path):
    records = get_rotulo_carpeta_records(desde, hasta)
    if not records:
        raise ValueError("No se encontraron registros para generar el Rótulo de Carpeta.")

    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(7)

    for i in range(0, len(records), 2):
        if i > 0:
            doc.add_page_break()

        _add_rotulo_carpeta_word_label(doc, records[i])

        if i + 1 < len(records):
            _add_rotulo_carpeta_word_label(doc, records[i + 1])

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)

    log_action(current_username, "generate_rotulo_carpeta_word", "rotulo_carpeta", details=str(out))
    return str(out)



def generate_rotulo_carpeta_word(current_username, desde, hasta, output_path):
    records = get_rotulo_carpeta_records(desde, hasta)
    if not records:
        raise ValueError("No se encontraron registros para generar el Rótulo de Carpeta.")

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(7)

    for i in range(0, len(records), 2):
        if i > 0:
            doc.add_page_break()

        _add_rotulo_carpeta_word_label(doc, records[i])

        if i + 1 < len(records):
            _add_rotulo_carpeta_word_label(doc, records[i + 1])

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)

    log_action(current_username, "generate_rotulo_carpeta_word", "rotulo_carpeta", details=str(out))
    return str(out)


def generate_rotulo_carpeta(current_username, desde, hasta, output_format, output_dir):
    if output_dir is None or str(output_dir).strip() == "":
        output_dir = str(PROJECT_ROOT / "salidas")

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"ROTULO_CARPETA_{timestamp}"

    if output_format == "Excel":
        out = output_dir / f"{base_name}.xlsx"
        return generate_rotulo_carpeta_excel(current_username, desde, hasta, out)

    if output_format == "PDF":
        xlsx_path = output_dir / f"{base_name}.xlsx"
        generate_rotulo_carpeta_excel(current_username, desde, hasta, xlsx_path)
        return generate_rotulo_carpeta_pdf_from_excel(xlsx_path, output_dir)

    if output_format == "Word":
        out = output_dir / f"{base_name}.docx"
        return generate_rotulo_carpeta_word(current_username, desde, hasta, out)

    raise ValueError("Formato de salida no soportado.")



def get_rotulo_caja_config():
    with db_cursor() as (_, cur):
        cur.execute("SELECT config_key, config_value FROM rotulo_caja_config ORDER BY config_key")
        data = {r["config_key"]: r["config_value"] for r in cur.fetchall()}

    for key, value in ROTULO_CAJA_CONFIG_DEFAULTS:
        data.setdefault(key, value)

    return data


def save_rotulo_caja_config(current_username, data: dict):
    with db_cursor() as (_, cur):
        for key, value in data.items():
            cur.execute(
                """
                INSERT INTO rotulo_caja_config(config_key, config_value)
                VALUES (?, ?)
                ON CONFLICT(config_key) DO UPDATE SET config_value = excluded.config_value
                """,
                (key, value or ""),
            )

    log_action(current_username, "save_rotulo_caja_config", "rotulo_caja")


def get_rotulo_caja_mapping():
    with db_cursor() as (_, cur):
        cur.execute("SELECT * FROM rotulo_caja_mapping ORDER BY display_order, id")
        rows = [dict(r) for r in cur.fetchall()]

    existing = {r["rotulo_field"] for r in rows}
    if len(existing) < len(ROTULO_CAJA_FIELDS):
        with db_cursor() as (_, cur):
            for idx, f in enumerate(ROTULO_CAJA_FIELDS, start=1):
                if f not in existing:
                    cur.execute(
                        """
                        INSERT INTO rotulo_caja_mapping(rotulo_field, mapping_type, mapping_value, display_order)
                        VALUES (?, 'field', '', ?)
                        """,
                        (f, idx),
                    )
        return get_rotulo_caja_mapping()

    return rows


def update_rotulo_caja_mapping(current_username, mapping_id, mapping_type, mapping_value):
    with db_cursor() as (_, cur):
        cur.execute(
            "UPDATE rotulo_caja_mapping SET mapping_type = ?, mapping_value = ? WHERE id = ?",
            (mapping_type, mapping_value or "", mapping_id),
        )

    log_action(current_username, "update_rotulo_caja_mapping", "rotulo_caja", record_id=str(mapping_id))


def _select_rotulo_caja_records(caja, desde, hasta):
    cfg = get_rotulo_caja_config()
    campo_busqueda = cfg.get("campo_busqueda", "Número de orden")
    campo_caja = cfg.get("campo_caja", "Caja")
    rows = list_records()

    caja_s = str(caja or "").strip()
    desde_s = str(desde or "").strip()
    hasta_s = str(hasta or desde_s).strip()

    filtered = []
    for row in rows:
        if caja_s:
            val_caja = str(_find_record_value_ci(row, campo_caja)).strip()
            if val_caja != caja_s:
                continue
        filtered.append(row)

    if not desde_s and not hasta_s:
        selected = filtered
    else:
        selected = []
        if not hasta_s:
            hasta_s = desde_s

        numeric_range = _is_number_like(desde_s) and _is_number_like(hasta_s)

        if numeric_range:
            ini = float(desde_s.replace(",", "."))
            fin = float(hasta_s.replace(",", "."))
            if ini > fin:
                ini, fin = fin, ini

            for row in filtered:
                raw = _find_record_value_ci(row, campo_busqueda)
                if _is_number_like(raw):
                    val = float(str(raw).strip().replace(",", "."))
                    if ini <= val <= fin:
                        selected.append(row)

            selected.sort(key=lambda r: float(str(_find_record_value_ci(r, campo_busqueda)).strip().replace(",", ".")) if _is_number_like(_find_record_value_ci(r, campo_busqueda)) else 0)
        else:
            allowed = {desde_s, hasta_s}
            for row in filtered:
                raw = str(_find_record_value_ci(row, campo_busqueda)).strip()
                if raw in allowed:
                    selected.append(row)

    return selected


def _extract_document_tokens(text):
    """
    Extrae cadenas tipo 2025-00082 desde el texto de unidad documental.
    """
    return re.findall(r"\b\d{4}-\d+\b", str(text or ""))


def _get_year_from_date_or_token(value):
    value = str(value or "").strip()
    if not value:
        return ""
    m = re.search(r"\b(\d{4})[-/]", value)
    if m:
        return m.group(1)
    m = re.search(r"\b(\d{4})\b", value)
    return m.group(1) if m else ""


def _get_rotulo_caja_values(caja, desde, hasta):
    cfg = get_rotulo_caja_config()
    records = _select_rotulo_caja_records(caja, desde, hasta)
    if not records:
        raise ValueError("No se encontraron registros para generar el Rótulo de Caja.")

    campo_unidad = cfg.get("campo_unidad_documental", "Nombre unidad documental")
    campo_fecha_inicial = cfg.get("campo_fecha_inicial", "Fecha inicial")
    campo_fecha_final = cfg.get("campo_fecha_final", "Fecha final")

    first_record = records[0]
    last_record = records[-1]

    first_unidad = str(_find_record_value_ci(first_record, campo_unidad)).strip()
    last_unidad = str(_find_record_value_ci(last_record, campo_unidad)).strip()

    first_tokens = _extract_document_tokens(first_unidad)
    last_tokens = _extract_document_tokens(last_unidad)

    consecutivo_doc = first_tokens[0] if first_tokens else first_unidad
    correlativo_doc = last_tokens[-1] if last_tokens else last_unidad

    detail_rows = []
    all_years = []

    for idx, rec in enumerate(records[:7], start=1):
        unidad = str(_find_record_value_ci(rec, campo_unidad)).strip()
        fecha_ini = str(_find_record_value_ci(rec, campo_fecha_inicial)).strip()
        fecha_fin = str(_find_record_value_ci(rec, campo_fecha_final)).strip()

        for source in [unidad, fecha_ini, fecha_fin]:
            for y in re.findall(r"\b(20\d{2}|19\d{2})\b", str(source or "")):
                try:
                    all_years.append(int(y))
                except Exception:
                    pass

        detail_rows.append({
            "no": idx,
            "asunto": unidad,
            "fecha_inicial": fecha_ini,
            "fecha_final": fecha_fin,
        })

    for rec in records[7:]:
        unidad = str(_find_record_value_ci(rec, campo_unidad)).strip()
        fecha_ini = str(_find_record_value_ci(rec, campo_fecha_inicial)).strip()
        fecha_fin = str(_find_record_value_ci(rec, campo_fecha_final)).strip()
        for source in [unidad, fecha_ini, fecha_fin]:
            for y in re.findall(r"\b(20\d{2}|19\d{2})\b", str(source or "")):
                try:
                    all_years.append(int(y))
                except Exception:
                    pass

    while len(detail_rows) < 7:
        detail_rows.append({
            "no": "",
            "asunto": "",
            "fecha_inicial": "",
            "fecha_final": "",
        })

    if all_years:
        anio_1 = str(min(all_years))
        anio_2 = str(max(all_years))
    else:
        anio_1 = ""
        anio_2 = ""

    return {
        "titulo": cfg.get("titulo", ""),
        "dependencia": cfg.get("dependencia", ""),
        "serie": cfg.get("serie", ""),
        "subserie": cfg.get("subserie", ""),
        "consecutivo": f"{cfg.get('texto_consecutivo', '').strip()} {consecutivo_doc}".strip(),
        "correlativo": f"{cfg.get('texto_correlativo', '').strip()} {correlativo_doc}".strip(),
        "anio_1": anio_1,
        "anio_2": anio_2,
        "caja": str(caja or "").strip(),
        "observaciones": cfg.get("observaciones", ""),
        "detail_rows": detail_rows,
        "records": records,
    }


def _apply_rotulo_caja_page_setup(ws, total_rows=31):
    ws.page_setup.orientation = "portrait"
    ws.page_setup.paperSize = ws.PAPERSIZE_LETTER
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.sheet_properties.pageSetUpPr = PageSetupProperties(fitToPage=True)
    ws.page_margins.left = 0.2
    ws.page_margins.right = 0.25
    ws.page_margins.top = 0.2
    ws.page_margins.bottom = 0.25
    ws.page_margins.header = 0.1
    ws.page_margins.footer = 0.1
    ws.print_area = f"A1:K{total_rows}"


def _build_single_rotulo_caja(ws, values, start_row):
    thin = Side(style="thin", color="000000")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    gray = PatternFill("solid", fgColor="EDEDED")

    for row in range(start_row, start_row + 28):
        ws.row_dimensions[row].height = 16

    ws.row_dimensions[start_row + 1].height = 17
    ws.row_dimensions[start_row + 2].height = 17
    ws.row_dimensions[start_row + 3].height = 17
    ws.row_dimensions[start_row + 4].height = 17
    ws.row_dimensions[start_row + 8].height = 22
    ws.row_dimensions[start_row + 9].height = 22
    for r in range(start_row + 15, start_row + 22):
        ws.row_dimensions[r].height = 38
    ws.row_dimensions[start_row + 24].height = 28
    ws.row_dimensions[start_row + 25].height = 26

    # Marco exterior guía de corte: columna A a J, fila 2 a 29.
    _apply_outer_cut_border(ws, start_row, start_row + 28, 1, 10)

    # Formato interno B:I.
    c0 = 2

    # Encabezado: logo + título/dependencia
    _merge(ws, start_row + 1, c0, start_row + 4, c0 + 1)
    if LOGO_PATH.exists():
        try:
            img = XLImage(str(LOGO_PATH))
            img.width = 78
            img.height = 56
            ws.add_image(img, f"C{start_row + 2}")
        except Exception as exc:
            print("No se pudo insertar el escudo:", exc)

    _merge(ws, start_row + 1, c0 + 2, start_row + 2, c0 + 7)
    ws.cell(row=start_row + 1, column=c0 + 2).value = values.get("titulo", "")
    _style_cell(ws.cell(row=start_row + 1, column=c0 + 2), border=None, bold=True, align="center", wrap=True)

    _merge(ws, start_row + 3, c0 + 2, start_row + 4, c0 + 7)
    ws.cell(row=start_row + 3, column=c0 + 2).value = values.get("dependencia", "")
    _style_cell(ws.cell(row=start_row + 3, column=c0 + 2), border=None, bold=True, align="center", wrap=True)

    # Borde externo completo del encabezado, sin línea horizontal interna.
    for row in range(start_row + 1, start_row + 5):
        for col in range(c0, c0 + 8):
            ws.cell(row=row, column=col).border = Border()

    for col in range(c0, c0 + 8):
        ws.cell(row=start_row + 1, column=col).border = Border(top=thin)
        ws.cell(row=start_row + 4, column=col).border = Border(bottom=thin)

    for row in range(start_row + 1, start_row + 5):
        # Lado izquierdo externo
        old = ws.cell(row=row, column=c0).border
        ws.cell(row=row, column=c0).border = Border(left=thin, top=old.top, bottom=old.bottom)

        # Divisor vertical entre logo y texto
        old = ws.cell(row=row, column=c0 + 1).border
        ws.cell(row=row, column=c0 + 1).border = Border(right=thin, top=old.top, bottom=old.bottom)

        # Lado derecho externo
        old = ws.cell(row=row, column=c0 + 7).border
        ws.cell(row=row, column=c0 + 7).border = Border(right=thin, top=old.top, bottom=old.bottom)

    # Serie / Subserie
    for row, label, value in [
        (start_row + 5, "SERIE:", values.get("serie", "")),
        (start_row + 6, "SUBSERIE:", values.get("subserie", "")),
    ]:
        _merge(ws, row, c0, row, c0 + 1)
        _merge(ws, row, c0 + 2, row, c0 + 7)
        ws.cell(row=row, column=c0).value = label
        ws.cell(row=row, column=c0 + 2).value = value
        _style_cell(ws.cell(row=row, column=c0), border=border, bold=True, align="left", fill=gray)
        _style_cell(ws.cell(row=row, column=c0 + 2), border=border, align="left", wrap=True)
        for col in list(range(c0 + 1, c0 + 2)) + list(range(c0 + 3, c0 + 8)):
            ws.cell(row=row, column=col).border = border

    # Consecutivo / Correlativo
    for row, label, value in [
        (start_row + 8, "CONSECUTIVO:", values.get("consecutivo", "")),
        (start_row + 9, "CORRELATIVO:", values.get("correlativo", "")),
    ]:
        _merge(ws, row, c0, row, c0 + 1)
        _merge(ws, row, c0 + 2, row, c0 + 7)
        ws.cell(row=row, column=c0).value = label
        ws.cell(row=row, column=c0 + 2).value = value
        _style_cell(ws.cell(row=row, column=c0), border=border, bold=True, align="left", fill=gray)
        _style_cell(ws.cell(row=row, column=c0 + 2), border=border, align="left", wrap=True)
        for col in list(range(c0 + 1, c0 + 2)) + list(range(c0 + 3, c0 + 8)):
            ws.cell(row=row, column=col).border = border

    # Años y caja
    anio_row = start_row + 11
    _merge(ws, anio_row, c0, anio_row, c0 + 1)
    ws.cell(row=anio_row, column=c0).value = "AÑOS:"
    _style_cell(ws.cell(row=anio_row, column=c0), border=border, bold=True, align="left", fill=gray)

    # Año menor en D y año mayor en E.
    ws.cell(row=anio_row, column=c0 + 2).value = values.get("anio_1", "")
    ws.cell(row=anio_row, column=c0 + 3).value = values.get("anio_2", "")
    _style_cell(ws.cell(row=anio_row, column=c0 + 2), border=border, align="center")
    _style_cell(ws.cell(row=anio_row, column=c0 + 3), border=border, align="center")

    _merge(ws, anio_row, c0 + 4, anio_row, c0 + 5)
    ws.cell(row=anio_row, column=c0 + 4).value = "N° CAJA"
    _style_cell(ws.cell(row=anio_row, column=c0 + 4), border=border, bold=True, align="center", fill=gray)

    _merge(ws, anio_row, c0 + 6, anio_row, c0 + 7)
    ws.cell(row=anio_row, column=c0 + 6).value = values.get("caja", "")
    _style_cell(ws.cell(row=anio_row, column=c0 + 6), border=border, align="center")

    for col in range(c0, c0 + 8):
        ws.cell(row=anio_row, column=col).border = border

    # Tabla detalle
    header_row = start_row + 13
    ws.cell(row=header_row, column=c0).value = "No"
    _merge(ws, header_row, c0 + 1, header_row, c0 + 4)
    ws.cell(row=header_row, column=c0 + 1).value = "ASUNTO CONTENIDO EN CADA UNIDAD (CARPETA)"
    _merge(ws, header_row, c0 + 5, header_row, c0 + 7)
    ws.cell(row=header_row, column=c0 + 5).value = "FECHAS EXTREMAS"

    _style_cell(ws.cell(row=header_row, column=c0), border=border, bold=True, align="center", fill=gray)
    _style_cell(ws.cell(row=header_row, column=c0 + 1), border=border, bold=True, align="center", fill=gray)
    _style_cell(ws.cell(row=header_row, column=c0 + 5), border=border, bold=True, align="center", fill=gray)
    for col in range(c0, c0 + 8):
        ws.cell(row=header_row, column=col).border = border

    sub_header_row = start_row + 14
    _merge(ws, sub_header_row, c0 + 1, sub_header_row, c0 + 4)
    ws.cell(row=sub_header_row, column=c0 + 5).value = "Fecha inicial"
    _merge(ws, sub_header_row, c0 + 6, sub_header_row, c0 + 7)
    ws.cell(row=sub_header_row, column=c0 + 6).value = "Fecha final"
    for col in range(c0, c0 + 8):
        ws.cell(row=sub_header_row, column=col).border = border
        _style_cell(
            ws.cell(row=sub_header_row, column=col),
            border=border,
            bold=col in [c0 + 5, c0 + 6],
            align="center",
            fill=gray if col in [c0 + 5, c0 + 6] else None
        )

    data_start = start_row + 15
    for idx, row_data in enumerate(values.get("detail_rows", []), start=0):
        row = data_start + idx
        ws.cell(row=row, column=c0).value = row_data["no"]
        _merge(ws, row, c0 + 1, row, c0 + 4)
        ws.cell(row=row, column=c0 + 1).value = row_data["asunto"]
        ws.cell(row=row, column=c0 + 5).value = row_data["fecha_inicial"]
        _merge(ws, row, c0 + 6, row, c0 + 7)
        ws.cell(row=row, column=c0 + 6).value = row_data["fecha_final"]

        _style_cell(ws.cell(row=row, column=c0), border=border, align="center")
        _style_cell(ws.cell(row=row, column=c0 + 1), border=border, align="left", wrap=True)
        _style_cell(ws.cell(row=row, column=c0 + 5), border=border, align="center", wrap=True)
        _style_cell(ws.cell(row=row, column=c0 + 6), border=border, align="center", wrap=True)

        for col in range(c0, c0 + 8):
            ws.cell(row=row, column=col).border = border

    # Observaciones
    obs_row = start_row + 23
    _merge(ws, obs_row, c0, obs_row, c0 + 7)
    ws.cell(row=obs_row, column=c0).value = "OBSERVACIONES"
    _style_cell(ws.cell(row=obs_row, column=c0), border=border, bold=True, align="left", fill=gray)

    _merge(ws, obs_row + 1, c0, obs_row + 2, c0 + 7)
    ws.cell(row=obs_row + 1, column=c0).value = values.get("observaciones", "")
    _style_cell(ws.cell(row=obs_row + 1, column=c0), border=border, align="left", wrap=True)

    for row in range(obs_row, obs_row + 3):
        for col in range(c0, c0 + 8):
            ws.cell(row=row, column=col).border = border

def generate_rotulo_caja_excel(current_username, caja, desde, hasta, output_path):
    values = _get_rotulo_caja_values(caja, desde, hasta)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Rótulo Caja"

    widths = {"A": 8, "B": 22, "C": 22, "D": 22, "E": 22, "F": 16, "G": 16, "H": 16, "I": 4}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    _build_single_rotulo_caja(ws, values, 2)
    _apply_rotulo_caja_page_setup(ws, total_rows=31)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out)

    log_action(current_username, "generate_rotulo_caja_excel", "rotulo_caja", details=str(out))
    return str(out)


def generate_rotulo_caja_pdf_from_excel(excel_path, output_dir):
    return generate_fuid_pdf_from_excel(excel_path, output_dir)


def generate_rotulo_caja_word(current_username, caja, desde, hasta, output_path):
    values = _get_rotulo_caja_values(caja, desde, hasta)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(8)

    table = doc.add_table(rows=15, cols=4)
    table.style = "Table Grid"

    if LOGO_PATH.exists():
        try:
            p = table.cell(0, 0).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(LOGO_PATH), width=Inches(0.9))
        except Exception:
            pass

    table.cell(0, 1).merge(table.cell(1, 3))
    table.cell(0, 1).text = f"{values.get('titulo','')}\n{values.get('dependencia','')}"

    rows = [
        ("SERIE:", values.get("serie", "")),
        ("SUBSERIE:", values.get("subserie", "")),
        ("CONSECUTIVO:", values.get("consecutivo", "")),
        ("CORRELATIVO:", values.get("correlativo", "")),
        ("AÑOS / CAJA:", f"{values.get('anio_1','')} - {values.get('anio_2','')} / Caja {values.get('caja','')}"),
    ]

    r = 2
    for label, value in rows:
        table.cell(r, 0).text = label
        table.cell(r, 1).merge(table.cell(r, 3))
        table.cell(r, 1).text = str(value or "")
        r += 1

    doc.add_paragraph("")
    detail = doc.add_table(rows=1, cols=4)
    detail.style = "Table Grid"
    detail.rows[0].cells[0].text = "No"
    detail.rows[0].cells[1].text = "ASUNTO CONTENIDO EN CADA UNIDAD (CARPETA)"
    detail.rows[0].cells[2].text = "Fecha inicial"
    detail.rows[0].cells[3].text = "Fecha final"

    for row_data in values.get("detail_rows", []):
        row_cells = detail.add_row().cells
        row_cells[0].text = str(row_data["no"])
        row_cells[1].text = str(row_data["asunto"])
        row_cells[2].text = str(row_data["fecha_inicial"])
        row_cells[3].text = str(row_data["fecha_final"])

    doc.add_paragraph("OBSERVACIONES:")
    doc.add_paragraph(str(values.get("observaciones", "")))

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)

    log_action(current_username, "generate_rotulo_caja_word", "rotulo_caja", details=str(out))
    return str(out)


def generate_rotulo_caja(current_username, caja, desde, hasta, output_format, output_dir):
    if output_dir is None or str(output_dir).strip() == "":
        output_dir = str(PROJECT_ROOT / "salidas")

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"ROTULO_CAJA_{timestamp}"

    if output_format == "Excel":
        out = output_dir / f"{base_name}.xlsx"
        return generate_rotulo_caja_excel(current_username, caja, desde, hasta, out)

    if output_format == "PDF":
        xlsx_path = output_dir / f"{base_name}.xlsx"
        generate_rotulo_caja_excel(current_username, caja, desde, hasta, xlsx_path)
        return generate_rotulo_caja_pdf_from_excel(xlsx_path, output_dir)

    if output_format == "Word":
        out = output_dir / f"{base_name}.docx"
        return generate_rotulo_caja_word(current_username, caja, desde, hasta, out)

    raise ValueError("Formato de salida no soportado.")




def get_audit_rows(limit=200):
    with db_cursor() as (_, cur):
        cur.execute("SELECT * FROM audit_log ORDER BY id DESC LIMIT ?", (limit,))
        return [dict(r) for r in cur.fetchall()]


def get_dashboard_stats():
    with db_cursor() as (_, cur):
        # Total records
        cur.execute("SELECT COUNT(*) as total FROM records")
        total_records = cur.fetchone()["total"]

        # Total users
        cur.execute("SELECT COUNT(*) as total FROM users")
        total_users = cur.fetchone()["total"]

        # Records created this week
        cur.execute("SELECT COUNT(*) as total FROM records WHERE created_at >= date('now', '-7 days')")
        records_week = cur.fetchone()["total"]

        # Most active user
        cur.execute("SELECT created_by, COUNT(*) as total FROM records GROUP BY created_by ORDER BY total DESC LIMIT 1")
        active_user_row = cur.fetchone()
        active_user = active_user_row["created_by"] if active_user_row else "N/A"

        # Recent activity (last 5)
        cur.execute("SELECT * FROM audit_log ORDER BY id DESC LIMIT 5")
        recent_activity = [dict(r) for r in cur.fetchall()]

        # Records by month for chart (last 6 months)
        cur.execute("""
            SELECT strftime('%Y-%m', created_at) as month, COUNT(*) as total 
            FROM records 
            WHERE created_at >= date('now', '-6 months')
            GROUP BY month 
            ORDER BY month ASC
        """)
        chart_data = [dict(r) for r in cur.fetchall()]

        return {
            "total_records": total_records,
            "total_users": total_users,
            "records_week": records_week,
            "active_user": active_user,
            "recent_activity": recent_activity,
            "chart_data": chart_data
        }
